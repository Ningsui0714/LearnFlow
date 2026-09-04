from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "golden-role-research-collaboration-handbook.md"
OUTPUT = ROOT / "docs" / "golden-role-open-research-collaboration-handbook.docx"

FONT_LATIN = "Arial Unicode MS"
FONT_CJK = "Arial Unicode MS"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203748"
MUTED = "667581"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
CAUTION = "7A5A00"
CAUTION_FILL = "FFF7DD"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=None, bold=None, italic=None, latin=FONT_LATIN, cjk=FONT_CJK):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)
    rfonts.set(qn("w:cs"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color="000000", bold=False):
    style.font.name = FONT_LATIN
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    rfonts.set(qn("w:cs"), FONT_CJK)


def set_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_left_border(paragraph, color=BLUE, size="18", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def set_repeat_control(paragraph, keep_with_next=False, keep_together=False, page_break_before=False):
    ppr = paragraph._p.get_or_add_pPr()
    for tag, enabled in (
        ("w:keepNext", keep_with_next),
        ("w:keepLines", keep_together),
        ("w:pageBreakBefore", page_break_before),
    ):
        existing = ppr.find(qn(tag))
        if enabled and existing is None:
            ppr.append(OxmlElement(tag))
        elif not enabled and existing is not None:
            ppr.remove(existing)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = paragraph.add_run("— ")
    set_run_font(lead, size=9, color=MUTED)
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])
    set_run_font(run, size=9, color=MUTED)
    tail = paragraph.add_run(" —")
    set_run_font(tail, size=9, color=MUTED)


def configure_page(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("ROLE ATLAS · 黄金岗位研究")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = p.add_run("\t开放式研究与协作手册")
    set_run_font(right, size=8.5, color=MUTED)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_field(fp)

    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, "1F2933", False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, BLUE, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, BLUE, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, DARK_BLUE, True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    if "Code Block" not in [style.name for style in doc.styles]:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    set_style_font(code, 9.5, INK, False)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.12
    code.paragraph_format.keep_together = True

    if "Lead Callout" not in [style.name for style in doc.styles]:
        callout = doc.styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = doc.styles["Lead Callout"]
    set_style_font(callout, 11.5, WHITE, True)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2


def add_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    rfonts.set(qn("w:cs"), FONT_CJK)
    rpr.append(rfonts)
    lvl.extend([start, num_fmt, lvl_text, suff, ppr, rpr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_inline_runs(paragraph, text, size=11, color="1F2933", bold=False):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size - 0.5, color=DARK_BLUE, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color, bold=bold)


def add_cover(doc, title):
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(105)
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("ROLE ATLAS · 团队研究讲义")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    set_run_font(run, size=28, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run("为广泛阅读、自由发现与阶段性共识提供共同语言")
    set_run_font(run, size=13, color=DARK_BLUE)

    callout = doc.add_paragraph(style="Lead Callout")
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(callout, BLUE)
    set_left_border(callout, color=DARK_BLUE, size="20", space="10")
    run = callout.add_run("开放入口 · 证据约束 · 允许反复 · 保留争议")
    set_run_font(run, size=11.5, color=WHITE, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(78)
    meta.paragraph_format.space_after = Pt(0)
    run = meta.add_run("适用于黄金岗位包研究团队｜2026 年 8 月")
    set_run_font(run, size=10, color=MUTED)
    meta.add_run().add_break(WD_BREAK.PAGE)


def add_code_block(doc, lines):
    if not lines:
        return
    for index, line in enumerate(lines):
        p = doc.add_paragraph(style="Code Block")
        set_paragraph_shading(p, LIGHT_GRAY)
        set_left_border(p, color=BLUE, size="16", space="8")
        if index == 0:
            p.paragraph_format.space_before = Pt(5)
        if index == len(lines) - 1:
            p.paragraph_format.space_after = Pt(7)
        run = p.add_run(line if line else " ")
        set_run_font(run, size=9.5, color=INK)


def add_markdown_body(doc, source_text, bullet_num_id):
    lines = source_text.splitlines()
    code_mode = False
    code_lines = []
    first_h1_skipped = False

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if code_mode:
                add_code_block(doc, code_lines)
                code_lines = []
                code_mode = False
            else:
                code_mode = True
            continue

        if code_mode:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line.startswith("# "):
            if not first_h1_skipped:
                first_h1_skipped = True
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, text, size=16, color=BLUE, bold=True)
            set_repeat_control(p, keep_with_next=True)
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph(style="Heading 2")
            if text == "注意":
                p.style = doc.styles["Heading 3"]
                p.clear()
                add_inline_runs(p, "注意事项", size=12, color=CAUTION, bold=True)
                set_repeat_control(p, keep_with_next=True)
            else:
                add_inline_runs(p, text, size=13, color=BLUE, bold=True)
                set_repeat_control(p, keep_with_next=True)
            continue

        if line.startswith("- "):
            p = doc.add_paragraph()
            apply_bullet(p, bullet_num_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            add_inline_runs(p, line[2:].strip())
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        add_inline_runs(p, line)


def main():
    source_text = SOURCE.read_text(encoding="utf-8")
    title_match = re.search(r"^#\s+(.+)$", source_text, re.MULTILINE)
    title = title_match.group(1) if title_match else "黄金岗位开放式研究与协作手册"

    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.subject = "Role Atlas 黄金岗位研究团队协作指南"
    doc.core_properties.author = "Role Atlas 研究团队"
    doc.core_properties.keywords = "黄金岗位, 岗位研究, 证据, 典型工作任务, 协作"
    doc.core_properties.comments = "由版本化 Markdown 研究手册生成"

    configure_page(doc)
    configure_styles(doc)
    bullet_num_id = add_bullet_numbering(doc)
    add_cover(doc, title)
    add_markdown_body(doc, source_text, bullet_num_id)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
